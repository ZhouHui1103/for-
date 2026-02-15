import os
import re
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_COLOR_INDEX

def parse_and_add_runs(paragraph, text):
    # Regex for tags
    # 1: **
    # 2: <font color="..."> (captures color value in group 3)
    # 4: </font>
    # 5: <mark>
    # 6: </mark>
    tag_pattern = re.compile(r'(\*\*)|(<font color="([^"]+)">)|(</font>)|(<mark>)|(</mark>)')
    
    pos = 0
    bold_stack = 0
    color_stack = [] 
    highlight_stack = 0
    
    while pos < len(text):
        match = tag_pattern.search(text, pos)
        if not match:
            remaining = text[pos:]
            if remaining:
                run = paragraph.add_run(remaining)
                if bold_stack > 0: run.bold = True
                if color_stack:
                    run.font.color.rgb = color_stack[-1]
                if highlight_stack > 0:
                    run.font.highlight_color = WD_COLOR_INDEX.TURQUOISE
            break
        
        start, end = match.span()
        
        # Text before tag
        if start > pos:
            segment = text[pos:start]
            run = paragraph.add_run(segment)
            if bold_stack > 0: run.bold = True
            if color_stack:
                run.font.color.rgb = color_stack[-1]
            if highlight_stack > 0:
                run.font.highlight_color = WD_COLOR_INDEX.TURQUOISE
        
        # Process tag
        tag_text = match.group(0)
        if tag_text == '**':
            if bold_stack > 0:
                bold_stack -= 1
            else:
                bold_stack += 1
        elif tag_text.startswith('<font'):
            color_name = match.group(3).lower() if match.group(3) else 'black'
            rgb = RGBColor(0, 0, 0)
            if color_name == 'red': rgb = RGBColor(255, 0, 0)
            elif color_name == 'blue': rgb = RGBColor(0, 0, 255)
            elif color_name == 'green': rgb = RGBColor(0, 128, 0)
            elif color_name == 'yellow': rgb = RGBColor(255, 215, 0)
            elif color_name == 'purple': rgb = RGBColor(128, 0, 128)
            
            color_stack.append(rgb)
        elif tag_text == '</font>':
            if color_stack: color_stack.pop()
        elif tag_text == '<mark>':
            highlight_stack += 1
        elif tag_text == '</mark>':
            if highlight_stack > 0: highlight_stack -= 1
            
        pos = end

def convert_file(filepath):
    doc = Document()
    
    # Default styling
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    # style.element.rPr.rFonts.set_hAnsi("Microsoft YaHei") # Removed to avoid error
    style.font.size = Pt(11)
    
    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    for line in lines:
        line = line.strip()
        if not line:
            # Add empty paragraph for spacing
            doc.add_paragraph()
            continue
            
        if line.startswith('### '):
            text = line[4:].strip()
            p = doc.add_heading(level=3)
            parse_and_add_runs(p, text)
        elif line.startswith('## '):
            text = line[3:].strip()
            p = doc.add_heading(level=2)
            parse_and_add_runs(p, text)
        elif line.startswith('# '):
            text = line[2:].strip()
            p = doc.add_heading(level=1)
            parse_and_add_runs(p, text)
        elif line.startswith('- '):
            text = line[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            parse_and_add_runs(p, text)
        elif line == '---':
            p = doc.add_paragraph('----------------------------------------')
            p.alignment = 1 # WD_ALIGN_PARAGRAPH.CENTER (1)
        else:
            p = doc.add_paragraph()
            parse_and_add_runs(p, line)
            
    output_path = filepath.replace('.md', '.docx')
    doc.save(output_path)
    print(f"Generated: {output_path}")

if __name__ == '__main__':
    files = [
        r'纯AI文案/钥钥/纯AI文稿-03-操盘手不开课.md'
    ]
    
    for f in files:
        if os.path.exists(f):
            convert_file(f)
        else:
            print(f"File not found: {f}")

